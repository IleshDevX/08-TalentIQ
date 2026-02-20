"""
TalentIQ — Engine 1: File Processing Engine
Extracts raw text from PDF and DOCX resume files.
Handles tables, text boxes, headers/footers — not just paragraphs.
"""

from __future__ import annotations

import logging
from xml.etree import ElementTree

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class FileProcessingEngine:
    """Extract raw text from uploaded resume files."""

    SUPPORTED = {".pdf", ".docx"}

    def extract_text(self, file_path: str) -> str:
        """
        Extract all readable text from a resume file.

        Parameters
        ----------
        file_path : str
            Path to the PDF or DOCX file.

        Returns
        -------
        str
            Extracted text content.

        Raises
        ------
        ValueError
            If the file format is not supported.
        RuntimeError
            If text extraction fails.
        """
        logger.info("Extracting text from %s", file_path)
        path_lower = file_path.lower()

        try:
            if path_lower.endswith(".pdf"):
                return self._extract_pdf(file_path)
            elif path_lower.endswith(".docx"):
                return self._extract_docx(file_path)
            else:
                raise ValueError(
                    f"Unsupported file format: {file_path}. "
                    f"Supported: {', '.join(self.SUPPORTED)}"
                )
        except (ValueError, FileNotFoundError):
            raise
        except Exception as exc:
            logger.exception("Text extraction failed for %s", file_path)
            raise RuntimeError(
                f"Failed to extract text from {file_path}: {exc}"
            ) from exc

    # ------------------------------------------------------------------

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        pages: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
                # Also try extracting from tables
                for table in (page.extract_tables() or []):
                    for row in (table or []):
                        cells = [str(c).strip() for c in (row or []) if c]
                        if cells:
                            pages.append(" | ".join(cells))
        result = "\n".join(pages)
        logger.info("PDF extracted: %d chars from %d page-blocks", len(result), len(pages))
        if not result.strip():
            logger.warning("PDF produced empty text — file may be scanned/image-only")
        return result

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """
        Extract text from DOCX including:
        - Regular paragraphs
        - Tables (cells)
        - Text boxes / shapes (via XML fallback)
        - Headers and footers
        """
        doc = Document(file_path)
        parts: list[str] = []

        # 1. Paragraphs (the standard approach)
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        # 2. Tables — many resumes are table-based layouts
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        # 3. Headers and footers (often contain name/contact info)
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if header_footer:
                    for p in header_footer.paragraphs:
                        text = p.text.strip()
                        if text:
                            parts.append(text)

        # 4. Text boxes / shapes — extract from raw XML
        #    Many designed resumes put content in text boxes
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
            "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
            "v": "urn:schemas-microsoft-com:vml",
        }
        try:
            body = doc.element.body
            xml_str = ElementTree.tostring(body, encoding="unicode")
            root = ElementTree.fromstring(xml_str)
            # Find <w:t> elements inside textboxes / shapes
            for t_elem in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                txt = (t_elem.text or "").strip()
                if txt and len(txt) > 1:
                    parts.append(txt)
        except Exception:
            pass  # XML fallback is best-effort

        # Deduplicate while preserving order (XML fallback may re-capture paragraphs)
        seen: set[str] = set()
        unique_parts: list[str] = []
        for part in parts:
            if part not in seen:
                unique_parts.append(part)
                seen.add(part)

        result = "\n".join(unique_parts)
        logger.info(
            "DOCX extracted: %d chars from %d text blocks "
            "(paragraphs=%d, tables=%d, sections=%d)",
            len(result),
            len(unique_parts),
            len(doc.paragraphs),
            len(doc.tables),
            len(doc.sections),
        )
        if not result.strip():
            logger.warning("DOCX produced empty text — file may have non-text content only")
        return result
